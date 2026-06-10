from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
SALIDAS_DIR = BASE_DIR / "salidas"
CSV_PATH = SALIDAS_DIR / "resumen_departamento_odpe.csv"
OUTPUT_PATH = SALIDAS_DIR / "ranking_odpe_por_eleccion.docx"


def hex_to_rgb(hex_color: str):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def rgb_to_hex(rgb):
    return "#%02x%02x%02x" % rgb


def blend(c1, c2, t):
    r1, g1, b1 = hex_to_rgb(c1)
    r2, g2, b2 = hex_to_rgb(c2)
    rgb = (
        round(r1 + (r2 - r1) * t),
        round(g1 + (g2 - g1) * t),
        round(b1 + (b2 - b1) * t),
    )
    return rgb_to_hex(rgb)


def ramp_color(value, vmin, vmax):
    low = "#d73027"
    mid = "#fdae61"
    high = "#1a9850"
    if pd.isna(value):
        return "#d9d9d9"
    if vmax <= vmin:
        return mid
    x = (value - vmin) / (vmax - vmin)
    x = max(0.0, min(1.0, x))
    if x <= 0.5:
        return blend(low, mid, x / 0.5)
    return blend(mid, high, (x - 0.5) / 0.5)


def font_color(bg_hex):
    r, g, b = hex_to_rgb(bg_hex)
    luminance = 0.299 * r + 0.587 * g + 0.114 * b
    return "FFFFFF" if luminance < 150 else "000000"


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.style = "Table Grid"
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell)


def prepare_data():
    if not CSV_PATH.exists():
        raise FileNotFoundError(f"No existe el CSV de entrada: {CSV_PATH}")

    df = pd.read_csv(CSV_PATH)
    df = df.rename(columns={"departameto": "departamento"})
    df["eleccion"] = df["eleccion"].astype(str)
    df["departamento"] = df["departamento"].astype(str)
    df["nombre_odpe"] = df["nombre_odpe"].astype(str)
    df["pct_mesas"] = pd.to_numeric(df["pct_mesas"], errors="coerce")
    df["pct_electores"] = pd.to_numeric(df["pct_electores"], errors="coerce")
    return df


def build_document(df):
    global_min_mesas = float(df["pct_mesas"].min())
    global_max_mesas = float(df["pct_mesas"].max())
    global_min_electores = float(df["pct_electores"].min())
    global_max_electores = float(df["pct_electores"].max())

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for margin in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, margin, Inches(0.5))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Ranking ODPE por avance")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Arial"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Ordenado de menor a mayor porcentaje de avance. "
        "El color rojo indica menor avance y el verde mayor avance."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Arial"

    doc.add_paragraph("")

    eleccion = df["eleccion"].dropna().iloc[0]
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading_run = heading.add_run("Elecci\u00f3n presidencial - segunda vuelta")
    heading_run.bold = True

    subset = df.sort_values(
        ["pct_mesas", "nombre_odpe", "odpe"],
        ascending=[True, True, True],
    ).reset_index(drop=True)
    subset.insert(0, "Rango", range(1, len(subset) + 1))
    subset["Avance mesas"] = subset["pct_mesas"].map(lambda x: f"{x:.1%}" if pd.notna(x) else "")
    subset["Avance electores"] = subset["pct_electores"].map(lambda x: f"{x:.1%}" if pd.notna(x) else "")

    headers = [
        "Rango",
        "Departamento",
        "ODPE",
        "Nombre ODPE",
        "Mesas proc.",
        "Mesas totales",
        "Avance mesas",
        "Electores proc.",
        "Electores totales",
        "Avance electores",
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=8)
        set_cell_shading(hdr_cells[i], "4F81BD")
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for _, row in subset.iterrows():
        row_cells = table.add_row().cells
        values = [
            row["Rango"],
            row["departamento"],
            row["odpe"],
            row["nombre_odpe"],
            int(row["mesas_procesadas"]) if pd.notna(row["mesas_procesadas"]) else "",
            int(row["mesas_totales"]) if pd.notna(row["mesas_totales"]) else "",
            row["Avance mesas"],
            int(row["electores_procesados"]) if pd.notna(row["electores_procesados"]) else "",
            int(row["electores_totales"]) if pd.notna(row["electores_totales"]) else "",
            row["Avance electores"],
        ]
        for j, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2, 4, 5, 6, 7, 8, 9] else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row_cells[j], value, size=7, align=align)
            set_cell_margins(row_cells[j])

        color_mesas = ramp_color(float(row["pct_mesas"]), global_min_mesas, global_max_mesas)
        color_electores = ramp_color(float(row["pct_electores"]), global_min_electores, global_max_electores)
        set_cell_shading(row_cells[6], color_mesas)
        set_cell_shading(row_cells[9], color_electores)
        for idx, bg in [(6, color_mesas), (9, color_electores)]:
            text_color = font_color(bg)
            for run in row_cells[idx].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(text_color)
                run.font.bold = True

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_run = note.add_run(f"Fuente: resumen_departamento_odpe.csv. Elecci\u00f3n procesada: {eleccion}.")
    note_run.font.size = Pt(8)
    note_run.font.name = "Arial"
    note_run.italic = True

    return doc


def main():
    df = prepare_data()
    doc = build_document(df)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Archivo generado: {OUTPUT_PATH}")
    print(f"Filas totales: {len(df)}")


if __name__ == "__main__":
    main()
